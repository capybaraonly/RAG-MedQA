
#

import json
import logging
import re
import csv
from copy import deepcopy
from io import BytesIO
from timeit import default_timer as timer
from openpyxl import load_workbook

from parser import get_text
from rag.nlp import is_english, random_choices, qbullets_category, add_positions, has_qbullet, docx_question_level
from rag.nlp import rag_tokenizer, tokenize_table, concat_img
from markdown import markdown

from common.float_utils import get_float


class Excel:
    def __call__(self, fnm, binary=None, callback=None):
        if not binary:
            wb = load_workbook(fnm)
        else:
            wb = load_workbook(BytesIO(binary))
        total = 0
        for sheetname in wb.sheetnames:
            total += len(list(wb[sheetname].rows))

        res, fails = [], []
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            rows = list(ws.rows)
            for i, r in enumerate(rows):
                q, a = "", ""
                for cell in r:
                    if not cell.value:
                        continue
                    if not q:
                        q = str(cell.value)
                    elif not a:
                        a = str(cell.value)
                    else:
                        break
                if q and a:
                    res.append((q, a))
                else:
                    fails.append(str(i + 1))
                if len(res) % 999 == 0:
                    callback(len(res) *
                             0.6 /
                             total, ("Extract pairs: {}".format(len(res)) +
                                     (f"{len(fails)} failure, line: %s..." %
                                      (",".join(fails[:3])) if fails else "")))

        callback(0.6, ("Extract pairs: {}. ".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        self.is_english = is_english(
            [rmPrefix(q) for q, _ in random_choices(res, k=30) if len(q) > 1])
        return res



class Docx:
    def __init__(self):
        pass

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        from docx import Document
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        qai_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ''
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6:  # not a question
                last_answer = f'{last_answer}\n{p_text}'
                current_image = self.get_picture(self.doc, p)
                last_image = concat_img(last_image, current_image)
            else:  # is a question
                if last_answer or last_image:
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        qai_list.append((sum_question, last_answer, last_image))
                    last_answer, last_image = '', None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = '\n'.join(question_stack)
            if sum_question:
                qai_list.append((sum_question, last_answer, last_image))

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return qai_list, tbls


def rmPrefix(txt):
    return re.sub(
        r"^(问题|答案|回答|user|assistant|Q|A|Question|Answer|问|答)[\t:： ]+", "", txt.strip(), flags=re.IGNORECASE)



def beAdocDocx(d, q, a, eng, image, row_num=-1):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if image:
        d["image"] = image
        d["doc_type_kwd"] = "image"
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def beAdoc(d, q, a, eng, row_num=-1):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def mdQuestionLevel(s):
    match = re.match(r'#*', s)
    return (len(match.group(0)), s.lstrip('#').lstrip()) if match else (0, s)


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
        Supported formats:
        - Excel (.xlsx/.xls): 2 columns, question then answer, no header
        - CSV/TXT: UTF-8, TAB or comma delimiter, 2 columns
        - JSON (.json): array of {"question":..., "answer":...} objects, or {"q":..., "a":...}
        - JSONL (.jsonl): one JSON object per line, same field conventions as JSON
        - PDF, Markdown, Docx: Q&A bullet/heading structure

        Every Q&A pair is treated as one atomic chunk.
        Malformed lines/records are skipped and reported.
    """
    eng = lang.lower() == "english"
    res = []
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    if re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = Excel()
        for ii, (q, a) in enumerate(excel_parser(filename, binary, callback)):
            res.append(beAdoc(deepcopy(doc), q, a, eng, ii))
        return res

    elif re.search(r"\.(txt)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        comma, tab = 0, 0
        for line in lines:
            if len(line.split(",")) == 2:
                comma += 1
            if len(line.split("\t")) == 2:
                tab += 1
        delimiter = "\t" if tab >= comma else ","

        fails = []
        question, answer = "", ""
        i = 0
        while i < len(lines):
            arr = lines[i].split(delimiter)
            if len(arr) != 2:
                if question:
                    answer += "\n" + lines[i]
                else:
                    fails.append(str(i + 1))
            elif len(arr) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = arr
            i += 1
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(lines)))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        return res

    elif re.search(r"\.(csv)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        delimiter = "\t" if any("\t" in line for line in lines) else ","

        fails = []
        question, answer = "", ""
        res = []
        reader = csv.reader(lines, delimiter=delimiter)

        for i, row in enumerate(reader):
            if len(row) != 2:
                if question:
                    answer += "\n" + lines[i]
                else:
                    fails.append(str(i + 1))
            elif len(row) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = row
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(list(reader))))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        return res

    elif re.search(r"\.(md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        _last_question, last_answer = "", ""
        question_stack, level_stack = [], []
        code_block = False
        for index, line in enumerate(lines):
            if line.strip().startswith('```'):
                code_block = not code_block
            question_level, question = 0, ''
            if not code_block:
                question_level, question = mdQuestionLevel(line)

            if not question_level or question_level > 6:  # not a question
                last_answer = f'{last_answer}\n{line}'
            else:  # is a question
                if last_answer.strip():
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        res.append(beAdoc(deepcopy(doc), sum_question,
                                          markdown(last_answer, extensions=['markdown.extensions.tables']), eng, index))
                    last_answer = ''

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(question)
                level_stack.append(question_level)
        if last_answer.strip():
            sum_question = '\n'.join(question_stack)
            if sum_question:
                res.append(beAdoc(deepcopy(doc), sum_question,
                                  markdown(last_answer, extensions=['markdown.extensions.tables']), eng, index))
        return res

    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        docx_parser = Docx()
        qai_list, tbls = docx_parser(filename, binary,
                                     from_page=0, to_page=10000, callback=callback)
        res = tokenize_table(tbls, doc, eng)
        for i, (q, a, image) in enumerate(qai_list):
            res.append(beAdocDocx(deepcopy(doc), q, a, eng, image, i))
        return res

    elif re.search(r"\.(json|jsonl)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.strip().splitlines() if re.search(r"\.jsonl$", filename, re.IGNORECASE) else None

        if lines is None:
            # JSON array format
            try:
                records = json.loads(txt)
                if isinstance(records, dict):
                    records = [records]
            except json.JSONDecodeError as ex:
                raise ValueError(f"Invalid JSON file: {ex}")
        else:
            # JSONL format — one record per line
            records = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    pass

        fails = []
        for i, record in enumerate(records):
            q = record.get("question") or record.get("q") or record.get("问题") or ""
            a = record.get("answer") or record.get("a") or record.get("回答") or record.get("答案") or ""
            q, a = str(q).strip(), str(a).strip()
            if q and a:
                res.append(beAdoc(deepcopy(doc), q, a, eng, i))
            else:
                fails.append(str(i + 1))
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / max(len(records), 1),
                         "Extract Q&A: {}".format(len(res)) +
                         (f", {len(fails)} failures" if fails else ""))

        callback(0.6, "Extract Q&A: {}".format(len(res)) +
                 (f", {len(fails)} failure(s) at record(s): {','.join(fails[:5])}" if fails else ""))
        return res

    raise NotImplementedError(
        "Supported formats: Excel, CSV/TXT, JSON, JSONL, PDF, Markdown, Docx.")


if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
